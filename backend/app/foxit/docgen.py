"""
DocForge CLI — Foxit Document Generation API client.

Uses the synchronous GenerateDocumentBase64 endpoint:
  POST {HOST}/document-generation/api/GenerateDocumentBase64

Important Foxit template engine rules (discovered empirically):
  1. Tokens require whitespace before the opening {{ — e.g. "v{{ x }}" fails,
     but "v {{ x }}" works.
  2. Tables sharing the same inner token names ({{ title }}, {{ description }})
     cause later tables to fail. Each table must use unique field names.
  3. Top-level field names should use camelCase to match Foxit conventions.
"""

import base64
import io
from typing import Any

import httpx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.foxit.auth import FoxitCredentials
from app.utils.logging import logger, step_timer


def _to_foxit_values(data: dict[str, Any]) -> dict[str, Any]:
    """
    Transform user-facing snake_case data into Foxit-compatible format.

    Top-level key renames (camelCase + avoid reserved words):
      product_name    → productName
      version         → releaseVersion
      release_date    → releaseDate
      breaking_changes→ breakingChanges  (also renames inner fields)

    Breaking changes inner field renames (avoid collision with features/fixes):
      title       → change
      description → detail
      migration   → migrate
    """
    top_map = {
        "version": "releaseVersion",
        "product_name": "productName",
        "release_date": "releaseDate",
        "breaking_changes": "breakingChanges",
    }
    out: dict[str, Any] = {}
    for key, value in data.items():
        new_key = top_map.get(key, key)

        if new_key == "fixes" and isinstance(value, list):
            value = [
                {
                    "id": item.get("id", ""),
                    "fixTitle": item.get("title", ""),
                    "fixDesc": item.get("description", ""),
                }
                for item in value
            ]

        if new_key == "breakingChanges" and isinstance(value, list):
            value = [
                {
                    "change": item.get("title", ""),
                    "detail": item.get("description", ""),
                    "migrate": item.get("migration", ""),
                }
                for item in value
            ]

        out[new_key] = value
    return out


def _build_docx_template(data: dict[str, Any]) -> bytes:
    """
    Programmatically build a Word document (.docx) using Foxit's token syntax.

    Rules applied:
      - Every {{ token }} has whitespace before the opening {{
      - Each table uses unique inner field names to avoid cross-table collision
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # --- Title ---
    title = doc.add_heading(level=0)
    run = title.add_run("{{ productName }}")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Subtitle with version (space before {{ is critical)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("Release Notes — Version {{ releaseVersion }}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = date_para.add_run("Release Date: {{ releaseDate }}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    # --- Summary ---
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("{{ summary }}")
    doc.add_paragraph("")

    # --- Features (inner tokens: title, description) ---
    doc.add_heading("New Features", level=1)
    if data.get("features"):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Feature"
        hdr[1].text = "Description"
        row = table.add_row().cells
        row[0].text = "{{TableStart:features}} {{ title }}"
        row[1].text = "{{ description }} {{TableEnd:features}}"
    else:
        doc.add_paragraph("None")
    doc.add_paragraph("")

    # --- Fixes (inner tokens: id, fixTitle, fixDesc — unique names) ---
    doc.add_heading("Bug Fixes", level=1)
    if data.get("fixes"):
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Title"
        hdr[2].text = "Description"
        row = table.add_row().cells
        row[0].text = "{{TableStart:fixes}} {{ id }}"
        row[1].text = "{{ fixTitle }}"
        row[2].text = "{{ fixDesc }} {{TableEnd:fixes}}"
    else:
        doc.add_paragraph("None")
    doc.add_paragraph("")

    # --- Breaking Changes (inner tokens: change, detail, migrate — unique names) ---
    h = doc.add_heading("Breaking Changes", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    if data.get("breaking_changes") or data.get("breakingChanges"):
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Change"
        hdr[1].text = "Description"
        hdr[2].text = "Migration"
        row = table.add_row().cells
        row[0].text = "{{TableStart:breakingChanges}} {{ change }}"
        row[1].text = "{{ detail }}"
        row[2].text = "{{ migrate }} {{TableEnd:breakingChanges}}"
    else:
        doc.add_paragraph("None")
    doc.add_paragraph("")

    # --- Links (inner tokens: label, url — unique to this table) ---
    doc.add_heading("Useful Links", level=1)
    if data.get("links"):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Label"
        hdr[1].text = "URL"
        row = table.add_row().cells
        row[0].text = "{{TableStart:links}} {{ label }}"
        row[1].text = "{{ url }} {{TableEnd:links}}"
    else:
        doc.add_paragraph("None")

    # --- Footer ---
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        "Generated by DocForge CLI — Powered by Foxit Document Generation API"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generate_pdf_from_template(
    base_url: str,
    credentials: FoxitCredentials,
    release_data: dict[str, Any],
) -> bytes:
    """
    Build a Word template from the release data, call the Foxit Document
    Generation API, and return the resulting PDF as raw bytes.

    Data keys are transformed for Foxit compatibility before sending.
    """
    with step_timer("Build Word template"):
        docx_bytes = _build_docx_template(release_data)
        b64_template = base64.b64encode(docx_bytes).decode("utf-8")

    doc_values = _to_foxit_values(release_data)

    payload = {
        "outputFormat": "pdf",
        "documentValues": doc_values,
        "base64FileString": b64_template,
    }

    endpoint = f"{base_url}/document-generation/api/GenerateDocumentBase64"

    with step_timer("Foxit Document Generation API call"):
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(
                endpoint,
                json=payload,
                headers=credentials.as_headers(),
            )
            resp.raise_for_status()
            result = resp.json()

    if result.get("base64FileString"):
        pdf_bytes = base64.b64decode(result["base64FileString"])
        logger.info(
            "  Document Generation returned %d bytes PDF", len(pdf_bytes)
        )
        return pdf_bytes

    raise RuntimeError(
        f"Document Generation API error: {result.get('message', 'unknown')}"
    )
